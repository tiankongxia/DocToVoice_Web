#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from __future__ import annotations

from email.parser import BytesParser
from email.policy import default as email_policy
import json
import os
import queue
import re
import shutil
import subprocess
import threading
import time
import uuid
from datetime import datetime
from http import HTTPStatus
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import quote, unquote, urlparse
from urllib.request import Request, urlopen

from docx import Document


BASE_DIR = Path(__file__).resolve().parent
STATIC_DIR = BASE_DIR / "static"
OUTPUT_DIR = BASE_DIR / "outputs"
TMP_DIR = BASE_DIR / "tmp"

OUTPUT_DIR.mkdir(exist_ok=True)
TMP_DIR.mkdir(exist_ok=True)

os.environ["PATH"] = "/opt/homebrew/bin:/usr/local/bin:/opt/homebrew/Caskroom/miniforge/base/bin:" + os.environ.get("PATH", "")
os.environ["FFMPEG_BINARY"] = "/opt/homebrew/bin/ffmpeg"
os.environ["FFPROBE_BINARY"] = "/opt/homebrew/bin/ffprobe"

FFMPEG = shutil.which("ffmpeg") or "/opt/homebrew/bin/ffmpeg"
EDGE_TTS = shutil.which("edge-tts") or "/opt/homebrew/Caskroom/miniforge/base/bin/edge-tts"

DEFAULT_VOICE = "zh-CN-YunjianNeural"
DEFAULT_RATE = "-5%"
DEFAULT_MAX_CHARS = 5000
DEFAULT_PAUSE_MS = 1000
TEMP_AUDIO_PREFIX = "temp_audio_"
MAX_OUTPUT_JOBS = 5
PUNCTUATION_TO_NEWLINE = "。，？！；、——!?：:…【】"
REMOVE_CHARS = "\"'""'' "

VOICES = [
    "zh-CN-YunjianNeural",
    "zh-CN-XiaoxiaoNeural",
    "zh-CN-YunxiNeural",
    "zh-CN-YunyangNeural",
    "zh-CN-XiaoyiNeural",
]

jobs: dict[str, "Job"] = {}


class MultipartForm:
    def __init__(self, fields: dict[str, str], files: dict[str, tuple[str, bytes]]):
        self.fields = fields
        self.files = files

    def getfirst(self, name: str, default: str = "") -> str:
        return self.fields.get(name, default)


class Job:
    def __init__(self, title: str):
        self.id = uuid.uuid4().hex
        self.title = title
        self.created_at = datetime.now().isoformat(timespec="seconds")
        self.status = "queued"
        self.progress = 0
        self.message = "已加入队列"
        self.files: list[dict[str, str]] = []
        self.error: str | None = None
        self.cancelled = False
        self.events: queue.Queue[dict] = queue.Queue()

    def cancel(self):
        self.cancelled = True
        self.emit(status="cancelled", message="已取消", progress=self.progress)

    def check_cancelled(self):
        if self.cancelled:
            raise RuntimeError("已取消")

    def emit(self, **payload):
        if "status" in payload:
            self.status = payload["status"]
        if "progress" in payload:
            self.progress = payload["progress"]
        if "message" in payload:
            self.message = payload["message"]
        if "error" in payload:
            self.error = payload["error"]
        if "files" in payload:
            self.files = payload["files"]
        event = {
            "id": self.id,
            "title": self.title,
            "status": self.status,
            "progress": self.progress,
            "message": self.message,
            "files": self.files,
            "error": self.error,
        }
        event.update(payload)
        self.events.put(event)


def safe_name(value: str, fallback: str = "朗读") -> str:
    value = value.strip() or fallback
    value = re.sub(r"[\\/:*?\"<>|\r\n]+", "_", value)
    return value[:60] or fallback


def read_docx(file_path: Path) -> list[str]:
    doc = Document(str(file_path))
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def clean_title_line(text: str) -> str:
    text = re.sub(r"\[\[(.*?)\|(.*?)\]\]", r"\1", text)
    return " ".join(text.replace("\u00a0", " ").split()).strip()


def suggested_title(paragraphs: list[str]) -> str:
    lines = [clean_title_line(para) for para in paragraphs]
    lines = [line for line in lines if line]
    if not lines:
        return "朗读"
    return "_".join(lines[:2])


def replace_titles_in_docx(filename: Path) -> Path:
    doc = Document(str(filename))
    pattern = r"《(.*?)》"
    replaced = False
    for para in doc.paragraphs:
        for run in para.runs:
            new = re.sub(pattern, r"[[《\1》|\1]]", run.text)
            if new != run.text:
                run.text = new
                replaced = True
    if not replaced:
        return filename
    new_name = filename.with_name(filename.stem + "_replaced.docx")
    doc.save(str(new_name))
    return new_name


def split_into_blocks(paragraphs: list[str], max_chars: int) -> list[list[str]]:
    blocks: list[list[str]] = []
    current: list[str] = []
    current_len = 0
    for para in paragraphs:
        if current and current_len + len(para) > max_chars:
            blocks.append(current)
            current = [para]
            current_len = len(para)
        else:
            current.append(para)
            current_len += len(para)
    if current:
        blocks.append(current)
    return blocks


def process_for_tts(text: str) -> str:
    return re.sub(r"\[\[(.*?)\|(.*?)\]\]", r"\2", text)


def process_for_subtitle(text: str) -> str:
    text = re.sub(r"\[\[(.*?)\|(.*?)\]\]", r"\1", text)
    title_map = {}

    def save_title(match):
        key = f"<TITLE_{len(title_map)}>"
        title_map[key] = match.group(0)
        return key

    text = re.sub(r"《.*?》", save_title, text)
    for punct in PUNCTUATION_TO_NEWLINE:
        text = text.replace(punct, "\n")
    text = re.sub(f"[{re.escape(REMOVE_CHARS)}]", "", text)
    for key, value in title_map.items():
        text = text.replace(key, value)
    return text.strip()


def save_subtitle(text: str, index: int, out_dir: Path) -> Path:
    path = out_dir / f"cc_{index + 1}.docx"
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(str(path))
    return path


def normalize_dropbox_url(url: str) -> str:
    if "dl=0" in url:
        return url.replace("dl=0", "dl=1")
    if "dl=1" not in url:
        return url + ("&" if "?" in url else "?") + "dl=1"
    return url


def download_docx(url: str, target: Path):
    req = Request(normalize_dropbox_url(url), headers={"User-Agent": "LangduWeb/1.0"})
    with urlopen(req, timeout=60) as response:
        target.write_bytes(response.read())


def as_int(value, default: int, minimum: int, maximum: int) -> int:
    try:
        parsed = int(value)
    except (TypeError, ValueError):
        return default
    return max(minimum, min(maximum, parsed))


def parse_multipart_form(handler: BaseHTTPRequestHandler) -> MultipartForm:
    content_type = handler.headers.get("Content-Type", "")
    content_length = int(handler.headers.get("Content-Length", "0"))
    body = handler.rfile.read(content_length)
    if "multipart/form-data" not in content_type:
        return MultipartForm({}, {})

    raw_message = (
        f"Content-Type: {content_type}\r\n"
        "MIME-Version: 1.0\r\n\r\n"
    ).encode("utf-8") + body
    message = BytesParser(policy=email_policy).parsebytes(raw_message)

    fields: dict[str, str] = {}
    files: dict[str, tuple[str, bytes]] = {}
    for part in message.iter_parts():
        name = part.get_param("name", header="content-disposition")
        if not name:
            continue
        filename = part.get_filename()
        payload = part.get_payload(decode=True) or b""
        if filename:
            files[name] = (filename, payload)
        else:
            charset = part.get_content_charset() or "utf-8"
            fields[name] = payload.decode(charset, errors="replace")
    return MultipartForm(fields, files)


def cleanup_old_outputs(keep: int = MAX_OUTPUT_JOBS):
    output_dirs = [item for item in OUTPUT_DIR.iterdir() if item.is_dir()]
    output_dirs.sort(key=lambda item: item.stat().st_mtime, reverse=True)
    for old_dir in output_dirs[keep:]:
        shutil.rmtree(old_dir, ignore_errors=True)


def list_audio_outputs() -> list[dict]:
    audio_files: list[dict] = []
    for job_dir in OUTPUT_DIR.iterdir():
        if not job_dir.is_dir():
            continue
        for mp3 in job_dir.glob("*.mp3"):
            stat = mp3.stat()
            audio_files.append(
                {
                    "jobId": job_dir.name,
                    "name": mp3.name,
                    "url": f"/outputs/{quote(job_dir.name)}/{quote(mp3.name)}",
                    "deleteUrl": f"/api/audio/{quote(job_dir.name)}/{quote(mp3.name)}",
                    "type": "audio",
                    "size": stat.st_size,
                    "modifiedAt": stat.st_mtime,
                }
            )
    audio_files.sort(key=lambda item: item["modifiedAt"], reverse=True)
    return audio_files


def delete_audio_output(job_id: str, filename: str) -> bool:
    target = (OUTPUT_DIR / job_id / filename).resolve()
    output_root = OUTPUT_DIR.resolve()
    try:
        target.relative_to(output_root)
    except ValueError:
        return False
    if target.suffix.lower() != ".mp3":
        return False
    if not target.exists():
        return False
    target.unlink()
    job_dir = target.parent
    if job_dir.exists() and not any(job_dir.glob("*.mp3")):
        shutil.rmtree(job_dir, ignore_errors=True)
    return True


def ffconcat_line(path: Path) -> str:
    escaped = str(path.resolve()).replace("'", "'\\''")
    return f"file '{escaped}'\n"


def make_silence_file(path: Path, pause_ms: int):
    if pause_ms <= 0 or path.exists():
        return
    duration = max(0.001, pause_ms / 1000)
    cmd = [
        FFMPEG,
        "-hide_banner",
        "-loglevel",
        "error",
        "-f",
        "lavfi",
        "-i",
        "anullsrc=r=24000:cl=mono",
        "-t",
        f"{duration:.3f}",
        "-q:a",
        "9",
        "-acodec",
        "libmp3lame",
        str(path),
    ]
    subprocess.run(cmd, check=True)


def concat_audio_files(parts: list[Path], target: Path):
    if not parts:
        raise RuntimeError("没有成功生成任何音频")
    concat_path = target.parent / f"{target.stem}_concat.txt"
    concat_path.write_text("".join(ffconcat_line(part) for part in parts), encoding="utf-8")
    try:
        cmd = [
            FFMPEG,
            "-y",
            "-hide_banner",
            "-loglevel",
            "error",
            "-f",
            "concat",
            "-safe",
            "0",
            "-i",
            str(concat_path),
            "-acodec",
            "libmp3lame",
            "-b:a",
            "192k",
            str(target),
        ]
        timeout_seconds = min(900, max(120, len(parts) * 8))
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=timeout_seconds,
            check=False,
        )
        if result.returncode != 0:
            detail = (result.stderr or result.stdout or "ffmpeg 合并失败").strip()
            raise RuntimeError(detail[-600:])
    except subprocess.TimeoutExpired as exc:
        raise RuntimeError(f"音频合并超时，已等待 {exc.timeout:.0f} 秒") from exc
    finally:
        concat_path.unlink(missing_ok=True)


def synthesize_block(
    block: list[str],
    target: Path,
    job: Job,
    voice: str,
    rate: str,
    pause_ms: int,
    total_paras: int,
    done_ref: list[int],
    block_label: str,
):
    temp_files: list[Path] = []
    concat_parts: list[Path] = []
    silence_file = target.parent / f"{target.stem}_silence.mp3"
    if pause_ms > 0:
        make_silence_file(silence_file, pause_ms)
        temp_files.append(silence_file)

    for index, para in enumerate(block):
        job.check_cancelled()
        temp_mp3 = target.parent / f"{TEMP_AUDIO_PREFIX}{target.stem}_{index}.mp3"
        tts_text = process_for_tts(para)
        ok = False

        for attempt in range(3):
            job.check_cancelled()
            job.emit(
                status="running",
                message=f"{block_label}：正在合成第 {index + 1}/{len(block)} 段",
            )
            cmd = [
                EDGE_TTS,
                "--text",
                tts_text,
                "--voice",
                voice,
                f"--rate={rate}",
                "--write-media",
                str(temp_mp3),
            ]
            subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, check=False)
            if temp_mp3.exists() and temp_mp3.stat().st_size > 0:
                ok = True
                break
            time.sleep(1 + attempt)
            job.check_cancelled()

        done_ref[0] += 1
        progress = int(done_ref[0] / max(1, total_paras) * 95)
        job.emit(progress=progress)

        if not ok:
            job.emit(message=f"{block_label}：第 {index + 1} 段失败，已跳过")
            continue

        temp_files.append(temp_mp3)
        concat_parts.append(temp_mp3)
        if pause_ms > 0:
            concat_parts.append(silence_file)

    try:
        job.check_cancelled()
        merge_progress = min(98, max(96, int(done_ref[0] / max(1, total_paras) * 98)))
        job.emit(progress=merge_progress, message=f"{block_label}：正在合并音频")
        concat_audio_files(concat_parts, target)
    finally:
        for item in temp_files:
            item.unlink(missing_ok=True)


def run_job(job: Job, payload: dict):
    work_dir = TMP_DIR / job.id
    out_dir = OUTPUT_DIR / job.id
    work_dir.mkdir(parents=True, exist_ok=True)
    out_dir.mkdir(parents=True, exist_ok=True)
    cleanup_old_outputs()

    try:
        job.check_cancelled()
        mode = payload["mode"]
        voice = payload["voice"] if payload["voice"] in VOICES else DEFAULT_VOICE
        rate = payload["rate"] or DEFAULT_RATE
        max_chars = payload["max_chars"]
        pause_ms = payload["pause_ms"]

        job.emit(status="running", progress=2, message="正在读取内容")
        job.check_cancelled()

        if mode in {"file", "url"}:
            docx_path = work_dir / "input.docx"
            if mode == "file":
                shutil.copyfile(payload["file_path"], docx_path)
                Path(payload["file_path"]).unlink(missing_ok=True)
            else:
                download_docx(payload["url"], docx_path)
            job.check_cancelled()
            docx_path = replace_titles_in_docx(docx_path)
            paragraphs = read_docx(docx_path)
            if not paragraphs:
                raise RuntimeError("文档里没有可朗读的内容")
            blocks = split_into_blocks(paragraphs, max_chars)
            title = safe_name(payload.get("title") or suggested_title(paragraphs))
            job.title = title
        else:
            raw = payload["text"].strip()
            paragraphs = [line.strip() for line in raw.splitlines() if line.strip()] or [raw]
            blocks = [paragraphs]
            title = safe_name(payload.get("title") or f"update-{datetime.now().strftime('%Y%m%d-%H%M%S')}")
            job.title = title

        total_paras = sum(len(block) for block in blocks)
        done_ref = [0]
        files: list[dict[str, str]] = []

        job.emit(progress=5, message=f"共 {len(blocks)} 个音频文件，{total_paras} 段")

        for block_index, block in enumerate(blocks):
            job.check_cancelled()
            if mode == "text":
                filename = f"{title}.mp3"
            elif len(blocks) == 1:
                filename = f"{title}.mp3"
            else:
                filename = f"{title}-{block_index + 1}.mp3"
            mp3_path = out_dir / filename
            label = f"音频 {block_index + 1}/{len(blocks)}"
            synthesize_block(block, mp3_path, job, voice, rate, pause_ms, total_paras, done_ref, label)
            files.append({"name": mp3_path.name, "url": f"/outputs/{job.id}/{mp3_path.name}", "type": "audio"})

            if mode in {"file", "url"}:
                subtitle_text = process_for_subtitle("\n".join(block))
                cc_path = save_subtitle(subtitle_text, block_index, out_dir)
                files.append({"name": cc_path.name, "url": f"/outputs/{job.id}/{cc_path.name}", "type": "docx"})

        job.emit(status="done", progress=100, message="完成", files=files)
    except Exception as exc:
        if job.cancelled:
            job.emit(status="cancelled", message="已取消")
        else:
            job.emit(status="error", error=str(exc), message="出错")
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)


class LangduHandler(BaseHTTPRequestHandler):
    server_version = "LangduWeb/1.0"

    def log_message(self, format, *args):
        print("[%s] %s" % (self.log_date_time_string(), format % args))

    def do_GET(self):
        parsed = urlparse(self.path)
        path = unquote(parsed.path)
        if path == "/":
            self.serve_file(STATIC_DIR / "index.html", "text/html; charset=utf-8")
        elif path == "/health":
            self.json_response({"ok": True})
        elif path.startswith("/static/"):
            self.serve_static(path)
        elif path.startswith("/outputs/"):
            self.serve_output(path)
        elif path == "/api/audio":
            self.json_response({"files": list_audio_outputs()})
        elif path.startswith("/api/jobs/") and path.endswith("/events"):
            self.serve_events(path)
        elif path.startswith("/api/jobs/"):
            self.serve_job(path)
        else:
            self.send_error(HTTPStatus.NOT_FOUND)

    def do_HEAD(self):
        parsed = urlparse(self.path)
        path = unquote(parsed.path)
        if path == "/":
            self.serve_file(STATIC_DIR / "index.html", "text/html; charset=utf-8", head_only=True)
        elif path.startswith("/static/"):
            self.serve_static(path, head_only=True)
        elif path.startswith("/outputs/"):
            self.serve_output(path, head_only=True)
        else:
            self.send_error(HTTPStatus.NOT_FOUND)

    def do_POST(self):
        if self.path == "/api/jobs":
            self.create_job()
        elif self.path.startswith("/api/jobs/") and self.path.endswith("/cancel"):
            self.cancel_job(self.path)
        else:
            self.send_error(HTTPStatus.NOT_FOUND)

    def do_DELETE(self):
        parsed = urlparse(self.path)
        path = unquote(parsed.path)
        if path.startswith("/api/audio/"):
            self.delete_audio(path)
        else:
            self.send_error(HTTPStatus.NOT_FOUND)

    def do_OPTIONS(self):
        self.send_response(HTTPStatus.NO_CONTENT)
        self.end_headers()

    def end_headers(self):
        self.send_cors_headers()
        super().end_headers()

    def send_cors_headers(self):
        self.send_header("Access-Control-Allow-Origin", os.environ.get("CORS_ALLOW_ORIGIN", "*"))
        self.send_header("Access-Control-Allow-Methods", "GET, POST, DELETE, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")

    def serve_static(self, path: str, head_only: bool = False):
        file_path = (STATIC_DIR / path.removeprefix("/static/")).resolve()
        if not str(file_path).startswith(str(STATIC_DIR.resolve())):
            self.send_error(HTTPStatus.FORBIDDEN)
            return
        content_type = "text/plain; charset=utf-8"
        if file_path.suffix == ".css":
            content_type = "text/css; charset=utf-8"
        elif file_path.suffix == ".js":
            content_type = "application/javascript; charset=utf-8"
        self.serve_file(file_path, content_type, head_only=head_only)

    def serve_output(self, path: str, head_only: bool = False):
        file_path = (BASE_DIR / path.removeprefix("/")).resolve()
        if not str(file_path).startswith(str(OUTPUT_DIR.resolve())):
            self.send_error(HTTPStatus.FORBIDDEN)
            return
        content_type = "application/octet-stream"
        if file_path.suffix == ".mp3":
            content_type = "audio/mpeg"
        self.serve_file(file_path, content_type, head_only=head_only)

    def serve_file(self, file_path: Path, content_type: str, head_only: bool = False):
        if not file_path.exists() or not file_path.is_file():
            self.send_error(HTTPStatus.NOT_FOUND)
            return
        data = file_path.read_bytes()
        self.send_response(HTTPStatus.OK)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        if not head_only:
            self.wfile.write(data)

    def serve_events(self, path: str):
        job_id = path.split("/")[3]
        job = jobs.get(job_id)
        if not job:
            self.send_error(HTTPStatus.NOT_FOUND)
            return

        self.send_response(HTTPStatus.OK)
        self.send_header("Content-Type", "text/event-stream; charset=utf-8")
        self.send_header("Cache-Control", "no-cache")
        self.send_header("Connection", "keep-alive")
        self.end_headers()

        job.emit()
        while True:
            try:
                event = job.events.get(timeout=20)
            except queue.Empty:
                event = {"status": job.status, "message": job.message, "progress": job.progress}
            data = "data: " + json.dumps(event, ensure_ascii=False) + "\n\n"
            try:
                self.wfile.write(data.encode("utf-8"))
                self.wfile.flush()
            except (BrokenPipeError, ConnectionResetError):
                break
            if event.get("status") in {"done", "error", "cancelled"}:
                break

    def serve_job(self, path: str):
        job_id = path.split("/")[3]
        job = jobs.get(job_id)
        if not job:
            self.send_error(HTTPStatus.NOT_FOUND)
            return
        self.json_response(
            {
                "id": job.id,
                "title": job.title,
                "status": job.status,
                "progress": job.progress,
                "message": job.message,
                "files": job.files,
                "error": job.error,
            }
        )

    def create_job(self):
        form = parse_multipart_form(self)

        mode = form.getfirst("mode", "text")
        raw_title = form.getfirst("title", "").strip()
        title = safe_name(raw_title, fallback="") if raw_title else ""
        voice = form.getfirst("voice", DEFAULT_VOICE)
        rate = form.getfirst("rate", DEFAULT_RATE)
        max_chars = as_int(form.getfirst("max_chars"), DEFAULT_MAX_CHARS, 500, 1000000)
        pause_ms = as_int(form.getfirst("pause_ms"), DEFAULT_PAUSE_MS, 0, 5000)

        payload = {
            "mode": mode,
            "title": title,
            "voice": voice,
            "rate": rate,
            "max_chars": max_chars,
            "pause_ms": pause_ms,
        }

        try:
            if mode == "file":
                file_item = form.files.get("file")
                if not file_item:
                    raise ValueError("请选择 .docx 文件")
                _filename, file_data = file_item
                upload_path = TMP_DIR / f"{uuid.uuid4().hex}.docx"
                upload_path.write_bytes(file_data)
                payload["file_path"] = upload_path
            elif mode == "url":
                url = form.getfirst("url", "").strip()
                if not url:
                    raise ValueError("请填写 Dropbox 文档链接")
                payload["url"] = url
            else:
                text = form.getfirst("text", "").strip()
                if not text:
                    raise ValueError("请输入要朗读的文字")
                payload["mode"] = "text"
                payload["text"] = text
        except Exception as exc:
            self.json_response({"error": str(exc)}, status=HTTPStatus.BAD_REQUEST)
            return

        job = Job(title=title)
        jobs[job.id] = job
        thread = threading.Thread(target=run_job, args=(job, payload), daemon=True)
        thread.start()
        self.json_response({"id": job.id})

    def delete_audio(self, path: str):
        parts = path.split("/", 4)
        if len(parts) != 5:
            self.send_error(HTTPStatus.NOT_FOUND)
            return
        _empty, _api, _audio, job_id, filename = parts
        if not delete_audio_output(job_id, filename):
            self.send_error(HTTPStatus.NOT_FOUND)
            return
        self.json_response({"ok": True, "files": list_audio_outputs()})

    def cancel_job(self, path: str):
        job_id = path.split("/")[3]
        job = jobs.get(job_id)
        if not job:
            self.send_error(HTTPStatus.NOT_FOUND)
            return
        job.cancel()
        self.json_response({"ok": True})

    def json_response(self, payload: dict, status=HTTPStatus.OK):
        data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        self.wfile.write(data)


def main():
    host = os.environ.get("LANGDU_HOST", "0.0.0.0")
    port = int(os.environ.get("PORT") or os.environ.get("LANGDU_PORT", "8765"))
    server = ThreadingHTTPServer((host, port), LangduHandler)
    print(f"朗读网页版已启动：http://localhost:{port}")
    print("同一 Wi-Fi 下，手机可访问这台电脑的局域网 IP 加端口。")
    server.serve_forever()


if __name__ == "__main__":
    main()
